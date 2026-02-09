#!/usr/bin/env python3
"""Extract ALL text content from a DOCX file - deep extraction including text boxes, shapes, etc."""

import zipfile
from lxml import etree

DOCX_PATH = "/Users/Paco/Documents/M2 GIP/INGENIE FISCALE DES PERSONNES PHYSIQUES & PATRIMOINE/Contrôle.docx"

# Namespaces used in OOXML
NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
    'v': 'urn:schemas-microsoft-com:vml',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

def extract_all_text_from_element(elem):
    """Extract ALL text nodes from any XML element, recursively."""
    texts = []
    for t in elem.iter():
        # Get w:t text nodes (main text)
        if t.tag == f'{{{NS["w"]}}}t' and t.text:
            texts.append(t.text)
        # Get a:t text nodes (drawing/shape text)
        elif t.tag == f'{{{NS["a"]}}}t' and t.text:
            texts.append(t.text)
        # VML text
        elif t.tag == f'{{{NS["v"]}}}textpath' and t.get('string'):
            texts.append(t.get('string'))
    return texts

def main():
    # First: standard python-docx extraction
    from docx import Document
    doc = Document(DOCX_PATH)

    print("=" * 80)
    print("STANDARD PARAGRAPHS (python-docx)")
    print("=" * 80)
    for i, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else "None"
        text = para.text
        print(f"[P{i:03d}] [Style: {style_name}] {text}")

    print("\n" + "=" * 80)
    print("TABLES (python-docx)")
    print("=" * 80)
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- Table {t_idx + 1} ({len(table.rows)} rows x {len(table.columns)} cols) ---")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text for cell in row.cells]
            print(f"  Row {r_idx}: {cells}")

    # Second: deep XML extraction from ZIP
    print("\n" + "=" * 80)
    print("DEEP XML EXTRACTION (all parts in the DOCX ZIP)")
    print("=" * 80)

    with zipfile.ZipFile(DOCX_PATH, 'r') as z:
        print(f"\nFiles in archive: {z.namelist()}")

        for name in z.namelist():
            if name.endswith('.xml') or name.endswith('.rels'):
                try:
                    data = z.read(name)
                    root = etree.fromstring(data)
                    texts = extract_all_text_from_element(root)
                    if texts:
                        print(f"\n--- {name} ---")
                        full_text = ''.join(texts)
                        print(full_text)
                except Exception as e:
                    print(f"\n--- {name} --- ERROR: {e}")

    # Third: paragraph-by-paragraph deep extraction from document.xml
    print("\n" + "=" * 80)
    print("PARAGRAPH-BY-PARAGRAPH FROM document.xml (including shapes/textboxes)")
    print("=" * 80)

    with zipfile.ZipFile(DOCX_PATH, 'r') as z:
        data = z.read('word/document.xml')
        root = etree.fromstring(data)

        # Find the body
        body = root.find(f'{{{NS["w"]}}}body')
        if body is None:
            print("No body found!")
            return

        # Iterate all children of body
        for idx, child in enumerate(body):
            tag = etree.QName(child).localname
            texts = extract_all_text_from_element(child)
            joined = ''.join(texts)
            if joined.strip():
                print(f"[Element {idx}] <{tag}> : {joined}")

    print("\n" + "=" * 80)
    print("EXTRACTION COMPLETE")
    print("=" * 80)

if __name__ == "__main__":
    main()
